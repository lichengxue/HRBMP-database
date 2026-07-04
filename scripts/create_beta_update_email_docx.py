from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_PATH = Path("docs/HRBMP_Database_GUI_Beta_1_1_2_Update_Email.docx")


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.10):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tbl_cell_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def next_numbering_id(numbering_root, tag_name, attr_name):
    values = []
    for node in numbering_root.findall(qn(tag_name)):
        value = node.get(qn(attr_name))
        if value is not None:
            values.append(int(value))
    return max(values, default=0) + 1


def create_decimal_numbering(doc, restart_at=1):
    numbering_root = doc.part.numbering_part.element
    abstract_id = next_numbering_id(numbering_root, "w:abstractNum", "w:abstractNumId")
    num_id = next_numbering_id(numbering_root, "w:num", "w:numId")

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), str(abstract_id))

    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract_num.append(multi_level)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")

    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)

    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl.append(num_fmt)

    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl.append(lvl_text)

    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    lvl.append(p_pr)

    abstract_num.append(lvl)
    numbering_root.append(abstract_num)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), str(restart_at))
    lvl_override.append(start_override)
    num.append(lvl_override)
    numbering_root.append(num)

    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)

    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")

    num_id_node = num_pr.find(qn("w:numId"))
    if num_id_node is None:
        num_id_node = OxmlElement("w:numId")
        num_pr.append(num_id_node)
    num_id_node.set(qn("w:val"), str(num_id))


def style_run(run, bold=False, size=11, color="000000"):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_label_paragraph(doc, label, value):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, after=6, line=1.10)
    label_run = paragraph.add_run(label)
    style_run(label_run, bold=True, size=11, color="0B2545")
    value_run = paragraph.add_run(value)
    style_run(value_run, size=11)
    return paragraph


def add_numbered_items(doc, items):
    for index, item in enumerate(items, start=1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        set_paragraph_spacing(paragraph, after=8, line=1.167)
        number = paragraph.add_run(f"{index}.\t")
        style_run(number, size=11)
        first = paragraph.add_run(item["lead"])
        style_run(first, bold=True, size=11, color="0B2545")
        rest = paragraph.add_run(item["body"])
        style_run(rest, size=11)


def build_docx():
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(title, after=4, line=1.10)
    title_run = title.add_run("HRBMP Database GUI Beta Version 1.1.2 Update")
    style_run(title_run, bold=True, size=18, color="0B2545")

    subtitle = doc.add_paragraph()
    set_paragraph_spacing(subtitle, after=14, line=1.10)
    subtitle_run = subtitle.add_run("Internal testing and feedback announcement")
    style_run(subtitle_run, size=11, color="555555")

    add_label_paragraph(doc, "Subject: ", "HRBMP Database GUI Beta Version 1.1.2 Released for Internal Testing and Feedback")

    greeting = doc.add_paragraph("Dear team,")
    set_paragraph_spacing(greeting, after=8, line=1.10)

    bluf_table = doc.add_table(rows=1, cols=1)
    bluf_table.autofit = False
    bluf_table.columns[0].width = Inches(6.5)
    bluf_cell = bluf_table.cell(0, 0)
    bluf_cell.width = Inches(6.5)
    set_cell_shading(bluf_cell, "F4F6F9")
    set_cell_margins(bluf_table)
    bluf_para = bluf_cell.paragraphs[0]
    set_paragraph_spacing(bluf_para, after=0, line=1.10)
    bluf_label = bluf_para.add_run("BLUF: ")
    style_run(bluf_label, bold=True, size=11, color="0B2545")
    bluf_text = bluf_para.add_run(
        "The HRBMP Database GUI beta version 1.1.2 has been released for internal testing and feedback. "
        "Please review the updated interface and share any comments, issues, or suggestions on content, navigation, layout, or usability."
    )
    style_run(bluf_text, size=11)

    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, after=4)

    doc.add_heading("Major updates", level=1)
    add_numbered_items(
        doc,
        [
            {
                "lead": "Photo Gallery added. ",
                "body": "The GUI now includes organized photo sections for Hudson River images, field sampling photos, sample warehouse images, and lab sample processing photos. Field and lab sampling images are grouped under clearer subheaders so the galleries are easier to browse.",
            },
            {
                "lead": "Educational Materials expanded. ",
                "body": "The Educational Materials section now includes K-12 Curriculum, Fish Tales species spotlights for 13 key species, History of HRBMP, Oral Interviews, Current Research/Publications, Classroom Materials, and Outreach Activities.",
            },
        ],
    )

    doc.add_heading("Minor updates", level=1)
    add_numbered_items(
        doc,
        [
            {
                "lead": "Geographic reference map updated. ",
                "body": "The geographic map on the home page was updated for clearer visual reference.",
            },
            {
                "lead": "Survey program descriptions added. ",
                "body": "Descriptions were added for the major HRBMP survey programs.",
            },
            {
                "lead": "Lab processing team added. ",
                "body": "A new team section was added for the lab processing team.",
            },
            {
                "lead": "News tab added. ",
                "body": "A News tab was added for future HRBMP program news and ongoing research updates.",
            },
            {
                "lead": "Search bar added. ",
                "body": "A search bar was added to the webpage tab ribbon to help users quickly find sections and subsections.",
            },
        ],
    )

    closing = doc.add_paragraph(
        "This is still an internal beta review version, so feedback on missing information, clarity, layout, navigation, and usability would be very helpful."
    )
    set_paragraph_spacing(closing, before=8, after=10, line=1.10)

    signoff = doc.add_paragraph("Best,")
    set_paragraph_spacing(signoff, after=2, line=1.10)
    signature = doc.add_paragraph()
    set_paragraph_spacing(signature, after=0, line=1.10)
    sig_run = signature.add_run("Database Development Team")
    style_run(sig_run, bold=True, size=11, color="0B2545")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("HRBMP Database GUI Beta Version 1.1.2")
    style_run(footer_run, size=9, color="555555")

    doc.save(OUT_PATH)
    print(OUT_PATH.resolve())


if __name__ == "__main__":
    build_docx()
